# import fitz  # PyMuPDF
#
# from docx import Document
# from docx.shared import Pt
# from docx.oxml.ns import qn
# from docx.oxml import OxmlElement
#
# # Open the PDF file
# pdf_path = "C:\\Users\\kondu\\PycharmProjects\\pythonProject\\main\\resources\\financial.pdf"
# pdf_document = fitz.open(pdf_path)
#
# # Create a Word document
# doc = Document()
#
# # Loop through each page
# for page_number in range(pdf_document.page_count):
#     page = pdf_document.load_page(page_number)
#     blocks = page.get_text("dict")["blocks"]  # Extract the text and style as a dictionary
#
#     for block in blocks:
#         if "lines" in block:
#             for line in block["lines"]:
#                 for span in line["spans"]:
#                     text = span["text"]
#                     font_name = span["font"]
#                     font_size = span["size"]
#
#                     # Create a paragraph in the Word document
#                     para = doc.add_paragraph()
#
#                     # Add the text with style
#                     run = para.add_run(text)
#                     run.font.name = font_name
#
#                     # Set the font size
#                     run.font.size = Pt(font_size)
#
#                     # Handle font style (bold/italic)
#                     if "Bold" in font_name:
#                         run.bold = True
#                     if "Italic" in font_name:
#                         run.italic = True
#
# # Save the Word document
# doc.save("output.docx")
#
# # Close the PDF document
# pdf_document.close()
#-----------------------------------------------------------------------------------------------------------------------------
# Updated Code
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open the PDF file
pdf_path = 'C:\\Users\\kondu\\PycharmProjects\\Innovation\\main\\resources\\smallPDF.pdf'
pdf_document = fitz.open(pdf_path)

# Create a Word document
doc = Document()

# Function to determine text alignment based on coordinates
def determine_alignment(span, page_width):
    bbox = span["bbox"]  # Get the bounding box of the text
    text_width = bbox[2] - bbox[0]
    left_margin = bbox[0]

    if left_margin < 10:  # Assuming small left margin means left-aligned
        return WD_ALIGN_PARAGRAPH.LEFT
    elif abs((page_width / 2) - (left_margin + text_width / 2)) < 10:  # Center-aligned
        return WD_ALIGN_PARAGRAPH.CENTER
    elif page_width - (left_margin + text_width) < 10:  # Right-aligned
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT  # Default to left alignment

# Loop through each page
for page_number in range(pdf_document.page_count):
    page = pdf_document.load_page(page_number)
    page_width = page.rect.width  # Get the page width
    blocks = page.get_text("dict")["blocks"]  # Extract the text and style as a dictionary

    for block in blocks:
        if "lines" in block:
            for line in block["lines"]:
                for span in line["spans"]:
                    text = span["text"]
                    font_name = span["font"]
                    font_size = span["size"]

                    # Create a paragraph in the Word document
                    para = doc.add_paragraph()

                    # Add the text with style
                    run = para.add_run(text)
                    run.font.name = font_name
                    run.font.size = Pt(font_size)

                    # Handle font style (bold/italic)
                    if "Bold" in font_name:
                        run.bold = True
                    if "Italic" in font_name:
                        run.italic = True

                    # Set the alignment of the paragraph based on the PDF text position
                    alignment = determine_alignment(span, page_width)
                    para.alignment = alignment

# Save the Word document
doc.save("output500.docx")

# Close the PDF document
pdf_document.close()
#-----------------------------------------------------------------------------------------------------------------------------
#Version 3
# import fitz  # PyMuPDF
# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.oxml import OxmlElement
# from docx.oxml.ns import qn
#
# # Open the PDF file
# pdf_path = "C:\\Users\\kondu\\PycharmProjects\\Innovation\\main\\resources\\somatosensory.pdf"
# # C:\Users\kondu\PycharmProjects\Innovation\main\resources\example.pdf
# pdf_document = fitz.open(pdf_path)
#
# # Create a Word document
# doc = Document()
#
# # Function to determine text alignment based on coordinates
# def determine_alignment(span, page_width):
#     bbox = span["bbox"]  # Get the bounding box of the text
#     text_width = bbox[2] - bbox[0]
#     left_margin = bbox[0]
#
#     if left_margin < 10:  # Assuming small left margin means left-aligned
#         return WD_ALIGN_PARAGRAPH.LEFT
#     elif abs((page_width / 2) - (left_margin + text_width / 2)) < 10:  # Center-aligned
#         return WD_ALIGN_PARAGRAPH.CENTER
#     elif page_width - (left_margin + text_width) < 10:  # Right-aligned
#         return WD_ALIGN_PARAGRAPH.RIGHT
#     return WD_ALIGN_PARAGRAPH.LEFT  # Default to left alignment
#
# # Function to check if a block might be a table (basic heuristic)
# def is_table(block):
#     lines = block.get("lines", [])
#     if len(lines) > 1:  # Tables usually have multiple rows
#         avg_line_height = sum(line["bbox"][3] - line["bbox"][1] for line in lines) / len(lines)
#         consistent_spacing = all(abs((lines[i+1]["bbox"][1] - lines[i]["bbox"][3])) < avg_line_height for i in range(len(lines)-1))
#         return consistent_spacing
#     return False

# Function to extract and write table data to the Word document
# def write_table_to_doc(block, doc):
#     rows = []
#     for line in block["lines"]:
#         row = []
#         for span in line["spans"]:
#             row.append(span["text"])
#         rows.append(row)
#
#     # Create table in the Word document
#     table = doc.add_table(rows=len(rows), cols=len(rows[0]))
#     table.style = 'Table Grid'
#
#     for i, row in enumerate(rows):
#         for j, cell_text in enumerate(row):
#             cell = table.cell(i, j)
#             cell.text = cell_text
#
# # Loop through each page
# for page_number in range(pdf_document.page_count):
#     page = pdf_document.load_page(page_number)
#     page_width = page.rect.width  # Get the page width
#     blocks = page.get_text("dict")["blocks"]  # Extract the text and style as a dictionary
#
#     for block in blocks:
#         if is_table(block):  # Check if the block might be a table
#             write_table_to_doc(block, doc)
#         else:
#             if "lines" in block:
#                 for line in block["lines"]:
#                     for span in line["spans"]:
#                         text = span["text"]
#                         font_name = span["font"]
#                         font_size = span["size"]
#
#                         # Create a paragraph in the Word document
#                         para = doc.add_paragraph()
#
#                         # Add the text with style
#                         run = para.add_run(text)
#                         run.font.name = font_name
#                         run.font.size = Pt(font_size)
#
#                         # Handle font style (bold/italic)
#                         if "Bold" in font_name:
#                             run.bold = True
#                         if "Italic" in font_name:
#                             run.italic = True
#
#                         # Set the alignment of the paragraph based on the PDF text position
#                         alignment = determine_alignment(span, page_width)
#                         para.alignment = alignment
#
# # Save the Word document
# doc.save("output_with_tables.docx")
#
# # Close the PDF document
# pdf_document.close()